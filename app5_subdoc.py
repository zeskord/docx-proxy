import base64
from docx.shared import Inches, Mm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


def app5_subdoc(data, doc):
    sd = doc.new_subdoc()
    file_entry = data["ОбмерныйПлан"]
    file_guid = ["ИмяФайла"]
    file_extension = file_entry["Расширение"]
    file_base64 = file_entry["ДанныеФайла"]
    image = base64.b64decode(file_base64)
    filename = f'temp/{file_guid}.{file_extension}'
    with open(filename, "wb") as file:
        file.write(image)

    par = sd.add_paragraph()
    par.add_run().add_picture(filename, width=None, height=Mm(200))
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    file_entry["Картинка"] = sd
    return sd