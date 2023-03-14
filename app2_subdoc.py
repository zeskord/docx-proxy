# Универсальный вывод картинки на весь лист.
import base64
from docx.shared import Inches, Mm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def picture_subdoc(doc, file_entry, max_width, max_height):
    sd = doc.new_subdoc()
    file_base64 = file_entry["ДанныеФайла"]
    image = base64.b64decode(file_base64)
    stream = io.BytesIO(image)
    par = sd.add_paragraph()
    # Небольшая возня с размерами.
    width = file_entry["Ширина"]
    height = file_entry["Высота"]
    k = width / height # Отношение сторон картинки.
    k2 = max_width / max_height
    if k < k2:
        # Это широкая картинка, ее нужно привести к максимальной ширине
        par.add_run().add_picture(stream, width=None, height=Mm(max_height))
    else:
        par.add_run().add_picture(stream, width=Mm(max_width), height=None)
    
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file_entry["Картинка"] = sd
    return sd