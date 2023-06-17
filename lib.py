import base64
import io
import os
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document
import uuid

# Добавляет картинку в прогон. Нужно передавать размер прямоугольника, в который требуется
# вписать картинку.
def picture(run, pic_data, max_width, max_height):
    file_base64 = pic_data["ДанныеФайла"]
    image = base64.b64decode(file_base64)
    stream = io.BytesIO(image)
    # Небольшая возня с размерами.
    width = pic_data["Ширина"]
    height = pic_data["Высота"]
    k = width / height # Отношение сторон картинки.
    k2 = max_width / max_height
    if k < k2:
        # Это широкая картинка, ее нужно привести к максимальной ширине
        run.add_picture(stream, width=None, height=Mm(max_height))
    else:
        run.add_picture(stream, width=Mm(max_width), height=None)

# Возращает поддокумент с переданной картинкой, вписанной в прямоугольник требуемого размера.
def picture_subdoc(doc, file_entry, max_width, max_height):
    sd = doc.new_subdoc()
    par = sd.add_paragraph()
    if file_entry["ИспользоватьПлейсхолдер"]:
        pass
    else:
        file_base64 = file_entry["ДанныеФайла"]
        image = base64.b64decode(file_base64)
        stream = io.BytesIO(image)
        # Небольшая возня с размерами.
        width = file_entry["Ширина"]
        height = file_entry["Высота"]
        k = width / height # Отношение сторон картинки.
        k2 = max_width / max_height
        if k < k2:
            # Это широкая картинка, ее нужно привести к максимальной ширине
            par.add_run().add_picture(stream, width=None, height=Mm(max_height))
        else:
            par.add_run().add_picture(stream, width=Mm(max_width), height=None)
    
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file_entry["Картинка"] = sd
    return sd


def table_tiled_pictures(doc, pics_data, max_width, max_height):
    table = doc.add_table(rows=0, cols=2)

    counter = 0
    pics_data_length = len(pics_data)
    for file_entry in pics_data:
        counter = counter + 1
        is_last_pic = counter == (pics_data_length)
        if counter % 2 != 0:
            row_pictures = table.add_row()
            # row_pictures.height = Mm(100)
            # row_descriptions = table.add_row()
        else:
            row_pictures = table.rows[len(table.rows) - 1]
            # row_descriptions = table.rows[len(table.rows) - 1]

        # Добавляем картинку в файл.
        current_col_id = 1 if counter % 2 == 0 else 0
        cell_pictures = row_pictures.cells[current_col_id]
        # Если это последняя картинка и она нечетная, то нужно объединить последние ячейки.
        if is_last_pic and current_col_id == 0:
            cell_pictures.merge(row_pictures.cells[current_col_id + 1])
        cell_pictures.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        picture_paragraph = cell_pictures.paragraphs[0]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_after = Pt(0)
        run = picture_paragraph.add_run()
        if file_entry["ИспользоватьПлейсхолдер"] == False:
            picture(run, file_entry, max_width, max_height)

        # cell_descriptions = row_descriptions.cells[current_col_id]
        # if is_last_pic and current_col_id == 0:
        #     cell_descriptions.merge(row_descriptions.cells[current_col_id + 1])
        # description = file_entry["Описание"]
        # par_description = cell_descriptions.paragraphs[0]
        # par_description.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # description_run = par_description.add_run(description)
        # description_run.font.size = Pt(12)

def temp_Document_for_tiled_pictures(pic_data):
    temp_doc = Document()
    table_tiled_pictures(temp_doc, pic_data["ДанныеКартинок"], pic_data["МаксимальнаяШирина"], pic_data["МаксимальнаяВысота"])
    return temp_doc

def tiled_pictures_subdoc(div_data, doc):
    # Создаем шаблон документа с нуля с помощью python-docx.
    temp_doc = temp_Document_for_tiled_pictures(div_data)
    temp_doc_filename = f"temp/{str(uuid.uuid4())}.docx"
    temp_doc.save(temp_doc_filename)


    final = doc.new_subdoc(temp_doc_filename)
    os.remove(temp_doc_filename)

    return final