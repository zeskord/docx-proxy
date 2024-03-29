
import json
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document

def volumes(data, doc):
    sd = doc.new_subdoc()
    return generate_table(data, sd)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def generate_table(data, doc):

    table = doc.add_table(rows=0, cols=0)
    table.style = "Table Grid"
    table.add_column(Mm(15))
    table.add_column(Mm(110))
    table.add_column(Mm(20))
    table.add_column(Mm(20))

    row_header = table.add_row()

    row_header.cells[0].paragraphs[0].add_run("№ п/п")
    row_header.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_header.cells[1].paragraphs[0].add_run("Наименование")
    row_header.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_header.cells[2].paragraphs[0].add_run("Ед.изм.")
    row_header.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_header.cells[3].paragraphs[0].add_run("Кол.")
    row_header.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные таблицы.
    table_data = data["ВедомостьВидовИОбъемовРабот"]["Данные"]
    # Объединения в 3 и 4 колонках.
    table_merges = data["ВедомостьВидовИОбъемовРабот"]["Объединения34"]
    for item in table_data:
        # subdoc = doc.new_subdoc()
        row = table.add_row()
        par = row.cells[0].paragraphs[0]
        run = par.add_run(item["Номер"])
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item["Наименование"] == "":
            run.bold = True
            row.cells[0].merge(row.cells[3])
        else:
            row.cells[1].paragraphs[0].add_run(item["Наименование"])
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].paragraphs[0].add_run(item["ЕдИзм"])
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].paragraphs[0].add_run(item["Количество"])
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for start, end in table_merges.items():
        start_row = int(start)
        start_cell = table.rows[start_row - 2].cells[2]
        end_cell = table.rows[end - 2].cells[3]
        start_cell.merge(end_cell)
        i = len(start_cell.paragraphs)
        while i > 0:
            i -= 1
            par = start_cell.paragraphs[i]
            if par.text == "":
                delete_paragraph(start_cell.paragraphs[i])


    return doc

if __name__ == '__main__':
    my_list = json.load(open("tests/volumes.json", encoding='utf-8-sig'))
    temp = {"Данные": my_list, "Объединения34": {"7": 8}}
    data = {"ВедомостьВидовИОбъемовРабот": temp}
    doc = Document()
    generate_table(data, doc)
    doc.save("temp/volumes.docx")  
