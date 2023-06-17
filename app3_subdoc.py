# Возвращает поддокумент приложения В.
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import lib

def app3_subdoc(data, doc):
    sd = doc.new_subdoc()
    
    max_width = data["КартинкиПриложенияВ_МаксимальныйРазмер"]["Ширина"]
    max_height = data["КартинкиПриложенияВ_МаксимальныйРазмер"]["Высота"]

    # Создаем стиль для примечаний к картинкам.
    picture_comment_style = sd.styles.add_style("Picture comment", WD_STYLE_TYPE.PARAGRAPH)
    picture_comment_style.font.italic = True
    picture_comment_style.font.size = Pt(8)

    picture_group_counter = 0
    for picture_group in data["КартинкиПриложенияВ"]:
        
        picture_group_counter += 1

        table = sd.add_table(rows=0, cols=2)

        counter = 0
        picture_group_length = len(picture_group["ДанныеФайлов"])
        for file_entry in picture_group["ДанныеФайлов"]:
            counter = counter + 1
            is_last_pic = counter == (picture_group_length)
            if counter % 2 != 0:
                row_pictures = table.add_row()
                # row_pictures.height = Mm(100)
                row_descriptions = table.add_row()
            else:
                row_pictures = table.rows[len(table.rows) - 2]
                row_descriptions = table.rows[len(table.rows) - 1]

            # Добавляем картинку в файл.
            current_col_id = 1 if counter % 2 == 0 else 0
            cell_pictures = row_pictures.cells[current_col_id]
            # Если это последняя картинка и она нечетная, то нужно объединить последние ячейки.
            if is_last_pic and current_col_id == 0:
                cell_pictures.merge(row_pictures.cells[current_col_id + 1])
            cell_pictures.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            picture_paragraph = cell_pictures.paragraphs[0]
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_paragraph.paragraph_format.space_after = Pt(0)
            run = picture_paragraph.add_run()
            if file_entry["ДанныеКартинки"]["ИспользоватьПлейсхолдер"] == False:
                lib.picture(run, file_entry["ДанныеКартинки"], max_width, max_height)

            cell_descriptions = row_descriptions.cells[current_col_id]
            if is_last_pic and current_col_id == 0:
                cell_descriptions.merge(row_descriptions.cells[current_col_id + 1])
            description = file_entry["Описание"]
            par_description = cell_descriptions.paragraphs[0]
            par_description.alignment = WD_ALIGN_PARAGRAPH.CENTER
            description_run = par_description.add_run(description)
            description_run.font.size = Pt(12)
        
        par = sd.add_paragraph(f'Рисунок {picture_group_counter}: {picture_group["ОписаниеРисунка"]}')
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(6)
        if "ПримечаниеРисунка" in picture_group and picture_group["ПримечаниеРисунка"] != "":
            par = sd.add_paragraph(f'Примечание - {picture_group["ПримечаниеРисунка"]}')
            par.style = picture_comment_style
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_after = Pt(6)
        
    return sd
